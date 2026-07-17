#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gera o curriculo final (Word) do Kevin Fernandes Silva - Red Team Junior."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x1F, 0x3B, 0x57)   # azul-escuro
GRAY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- margens e estilo base ----
for section in doc.sections:
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(2)
rpr = normal.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rpr.append(rFonts)
rFonts.set(qn('w:eastAsia'), 'Calibri')


def add_rule(paragraph):
    """Adiciona uma linha horizontal abaixo do paragrafo."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), '1F3B57')
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT
    add_rule(p)
    return p


def sub_heading(text, bold=True, size=10.5, color=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def body(text, space_after=4, size=10.5, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    return p


def bullet(text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def bullet_bold_lead(lead, rest, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(lead)
    r1.bold = True
    r1.font.size = Pt(size)
    r2 = p.add_run(rest)
    r2.font.size = Pt(size)
    return p


# ---------------------------------------------------------------------------
# CABECALHO
# ---------------------------------------------------------------------------
name_p = doc.add_paragraph()
name_p.paragraph_format.space_after = Pt(0)
name_run = name_p.add_run("Kevin Fernandes Silva")
name_run.bold = True
name_run.font.size = Pt(19)
name_run.font.color.rgb = ACCENT

role_p = doc.add_paragraph()
role_p.paragraph_format.space_after = Pt(2)
role_run = role_p.add_run("Desenvolvedor Fullstack | Aspirante a Red Team Júnior (Segurança Ofensiva)")
role_run.font.size = Pt(12)
role_run.font.color.rgb = GRAY

contact_p = doc.add_paragraph()
contact_p.paragraph_format.space_after = Pt(6)
contact_run = contact_p.add_run(
    "São Paulo, Brasil  ·  linkedin.com/in/kevin-fernandes-silva  ·  Inglês avançado"
)
contact_run.font.size = Pt(10)
contact_run.font.color.rgb = GRAY
add_rule(contact_p)

# ---------------------------------------------------------------------------
# RESUMO
# ---------------------------------------------------------------------------
section_heading("Resumo")
body(
    "Desenvolvedor Fullstack com experiência prática em C#, .NET, Angular, Golang, Python, "
    "Docker e bancos NoSQL (MongoDB, Redis), em transição ativa para Segurança Ofensiva com "
    "foco em Red Team Júnior. Uso recorrente de Nmap para reconhecimento de rede, conhecimento "
    "aplicado de TCP/IP, DNS, VPN e portas, e prática constante na plataforma TryHackMe "
    "complementam uma base técnica que vai além do ferramental: a vivência como desenvolvedor "
    "de sistemas conteinerizados e APIs permite leitura direta de superfícies de ataque reais — "
    "exposição de portas, autenticação, tráfego HTTP/HTTPS. Interesse particular em criptografia, "
    "automação ofensiva em Python e tópicos emergentes de segurança de IA."
)

# ---------------------------------------------------------------------------
# STACK TECNICA
# ---------------------------------------------------------------------------
section_heading("Stack técnica")

sub_heading("Segurança Ofensiva / Redes", color=ACCENT)
bullet_bold_lead("Nmap — ", "uso prático e recorrente para varredura e reconhecimento de rede")
bullet_bold_lead("Wireshark — ", "em desenvolvimento ativo; leitura e análise de pacotes")
bullet_bold_lead("Burp Suite — ", "uso básico (interceptação e manipulação de requisições HTTP)")
bullet_bold_lead("TCP/IP, DNS, VPN, portas e protocolos de rede — ", "conhecimento aplicado")
bullet_bold_lead("Linux — ", "uso avançado no dia a dia (administração, SSH, scripting)")
bullet_bold_lead("SSH — ", "acesso remoto, gestão de chaves e credenciais")
bullet_bold_lead("Criptografia clássica — ", "César, Vigenère, XOR, análise de frequência")
bullet_bold_lead("TryHackMe — ", "prática recorrente (reconhecimento, exploração web, fundamentos ofensivos)")

sub_heading("Desenvolvimento", color=ACCENT)
bullet("C#, .NET, Angular, Golang, Python, Docker, API Gateway, MongoDB, Redis")

sub_heading("Automação & Dados", color=ACCENT)
bullet_bold_lead("OCR aplicado (PaddleOCR) — ", "extração e automação de dados")
bullet_bold_lead("LLMs — ", "conhecimento introdutório e interesse em automação e segurança de IA")

# ---------------------------------------------------------------------------
# EXPERIENCIA PROFISSIONAL
# ---------------------------------------------------------------------------
section_heading("Experiência profissional")

sub_heading("HS Prevent — São Paulo, Brasil (1 ano e 5 meses)", size=11, color=ACCENT)

sub_heading("Desenvolvedor Fullstack | Tempo integral | set/2025 – atual (11 meses)", size=10.5, italic=True)
body(
    "Continuidade da atuação como estagiário, assumindo responsabilidade plena sobre manutenção "
    "e desenvolvimento de sistemas em C#, .NET, Angular, Golang e Python, incluindo componentes "
    "conteinerizados em Docker e integração com bancos NoSQL (MongoDB, Redis)."
)

sub_heading("Estagiário de TI | mar/2025 – set/2025 (7 meses)", size=10.5, italic=True)
bullet("Manutenção e desenvolvimento de sistemas em C#, .NET, Angular, Golang e Python")
bullet("Manutenção e implementação de conteinerização em Docker")
bullet("Acesso, análise e modificação de dados em bancos não relacionais (MongoDB, Redis)")
bullet("Testes funcionais seguindo metodologia definida")
bullet("Elaboração de documentos de critérios de aceitação para testes e requisitos de implantação, em apoio à equipe de DevOps")

# ---------------------------------------------------------------------------
# FORMACAO ACADEMICA
# ---------------------------------------------------------------------------
section_heading("Formação acadêmica")
sub_heading("Universidade Anhanguera São Paulo", size=11, color=ACCENT)
body("Análise e Desenvolvimento de Sistemas — fev/2023 – jun/2026 (em andamento)")

# ---------------------------------------------------------------------------
# SEGURANCA DA INFORMACAO - FORMACAO PRATICA E CTFS
# ---------------------------------------------------------------------------
section_heading("Segurança da Informação — Formação prática e CTFs")

sub_heading("TryHackMe — prática recorrente", color=ACCENT)
body(
    "Uso contínuo da plataforma para reforço prático de reconhecimento de rede, exploração web "
    "e fundamentos de Red Team, em complemento às trilhas estruturadas abaixo."
)

sub_heading("OverTheWire Bandit (níveis 0–12) — Linux, Bash e SSH", color=ACCENT)
body("Prática progressiva de administração de sistema via terminal, com autenticação remota real via SSH a cada nível.", space_after=2)
bullet("Acesso remoto e autenticação: ssh, login remoto, gestão de credenciais entre níveis")
bullet("Manipulação de arquivos: nomes com espaços/caracteres especiais, arquivos ocultos (ls -la), identificação de tipo de arquivo (file)")
bullet("Busca e filtragem: find por tamanho/permissão/dono/grupo, grep para busca textual por padrões")
bullet("Processamento de dados: sort, uniq, strings (extração de texto de binários)")
bullet("Encoding e compressão: Base64, ROT13, hexadecimal (xxd), compressão/descompressão encadeada (gzip, bzip2, tar)")

sub_heading("OverTheWire Krypton (níveis 0–5) — Criptografia clássica", color=ACCENT)
bullet("Diferença entre encoding e encriptação")
bullet("Cifra de César e ROT13 (força bruta e reconhecimento de padrão)")
bullet("Substituição monoalfabética com análise de frequência")
bullet("Cifra de Vigenère (criptografia polialfabética)")
bullet("Criptanálise estatística básica")

sub_heading("OverTheWire Natas (níveis 0–6) — Segurança Web", color=ACCENT)
bullet("Reconhecimento e enumeração web: view-source, robots.txt, descoberta de diretórios/arquivos ocultos")
bullet("HTTP na prática: inspeção e manipulação de headers (incluindo Referer) com curl, wget e Burp Suite")
bullet("Sessões e autenticação: análise de cookies, bypass de restrições client-side via DevTools")
bullet("Code review: leitura de código-fonte PHP, identificação de includes e arquivos de configuração expostos")

sub_heading("picoCTF — General Skills", color=ACCENT)
bullet("Linux/shell: grep, find, pipes, redirecionamento, cron, variáveis de ambiente, permissões")
bullet("Rede: netcat, requisições web, curl/wget, fundamentos de rede")
bullet("Dados: Git básico, hashes (md5, sha1, sha256), Base64, hex, ROT13, compressão, regex")
bullet("Python aplicado a scripts curtos de resolução de desafios")

sub_heading("cmd Challenge (níveis 0–20) — Terminal Linux avançado", color=ACCENT)
bullet("Arquivos e permissões: touch, cp, mv, rm, chmod, chown")
bullet("Texto e busca: grep, find, head, tail, cut, awk, sed, sort, uniq, wc, tr")
bullet("Pipes, redirecionamento e processos: |, <, tee, xargs, ps, kill, top")
bullet("Ambiente: PATH, variáveis de ambiente, wildcards, expansion, escaping, quoting")

sub_heading("Imersão Hacker (prof. Robson Costa)", color=ACCENT)
body("CVE e análise de vulnerabilidades, OSINT, Google Dork, tipos de malware, SEToolkit, Shodan, Wifite, investigação cibernética.")

# ---------------------------------------------------------------------------
# DIFERENCIAIS
# ---------------------------------------------------------------------------
section_heading("Diferenciais")
bullet("Combinação pouco comum entre desenvolvimento fullstack (C#, .NET, Angular, Golang, Python, Docker, API Gateway) e prática ofensiva hands-on (Nmap, Wireshark, TryHackMe, CTFs), permitindo leitura tanto do código quanto do tráfego de rede por trás de uma vulnerabilidade")
bullet("Interesse particular e aprofundado em criptografia clássica e aplicada")
bullet("Exposição prática a automação de dados (OCR com PaddleOCR) e a LLMs, incluindo tópicos emergentes de segurança de IA")
bullet("Perfil analítico e investigativo, validado por prática constante em TryHackMe, OverTheWire e picoCTF")
bullet("Inglês avançado")

OUT = "/tmp/claude-1000/-tmp-tmp-tVKEwguj8x/b05a02f7-697a-41c0-a65a-899ee4663f98/scratchpad/Kevin_Fernandes_Silva_CV_RedTeam.docx"
doc.save(OUT)
print("Salvo em:", OUT)
